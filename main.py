from __future__ import annotations

import os
import re
import shutil
import tempfile
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

from docx import Document
from PyPDF2 import PdfReader, PdfWriter


PLACEHOLDER_PATTERN = re.compile(r"\{\{([A-Za-z0-9_]+)\}\}")


def extract_placeholders(docx_path: str) -> list[str]:
    document = Document(docx_path)
    found = set()

    def collect_from_paragraphs(paragraphs):
        for paragraph in paragraphs:
            for match in PLACEHOLDER_PATTERN.findall(paragraph.text):
                found.add(match)

    collect_from_paragraphs(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                collect_from_paragraphs(cell.paragraphs)

    for section in document.sections:
        collect_from_paragraphs(section.header.paragraphs)
        collect_from_paragraphs(section.footer.paragraphs)

    return sorted(found)


def replace_placeholders(docx_path: str, replacements: dict[str, str], output_path: str) -> None:
    document = Document(docx_path)

    def replace_in_paragraphs(paragraphs):
        for paragraph in paragraphs:
            original_text = paragraph.text
            new_text = original_text
            for key, value in replacements.items():
                placeholder = f"{{{{{key}}}}}"
                new_text = new_text.replace(placeholder, value)
            if new_text != original_text:
                for run in list(paragraph.runs):
                    paragraph._p.remove(run._element)
                paragraph.add_run(new_text)

    replace_in_paragraphs(document.paragraphs)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)

    for section in document.sections:
        replace_in_paragraphs(section.header.paragraphs)
        replace_in_paragraphs(section.footer.paragraphs)

    document.save(output_path)


def read_docx_metadata(path: str) -> dict[str, str]:
    document = Document(path)
    props = document.core_properties
    return {
        "DC:TITLE": props.title or "",
        "DC:CREATOR": props.author or props.creator or "",
        "CP:DESCRIPTION": props.description or "",
        "CP:CATEGORY": props.category or "",
    }


def write_docx_metadata(path: str, title: str, creator: str, description: str, category: str) -> None:
    document = Document(path)
    props = document.core_properties
    props.title = title or None
    props.author = creator or None
    props.creator = creator or None
    props.description = description or None
    props.category = category or None
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        temp_name = tmp.name
    document.save(temp_name)
    shutil.move(temp_name, path)


def clear_docx_metadata(path: str) -> None:
    write_docx_metadata(path, "", "", "", "")


def read_pdf_metadata(path: str) -> dict[str, str]:
    reader = PdfReader(path)
    metadata = reader.metadata or {}
    return {
        "DC:TITLE": metadata.get("/Title", ""),
        "DC:CREATOR": metadata.get("/Author", ""),
        "CP:DESCRIPTION": metadata.get("/Subject", ""),
        "CP:CATEGORY": metadata.get("/Category", ""),
    }


def write_pdf_metadata(path: str, title: str, creator: str, description: str, category: str) -> None:
    reader = PdfReader(path)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.add_metadata({
        "/Title": title or "",
        "/Author": creator or "",
        "/Subject": description or "",
        "/Category": category or "",
    })
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        writer.write(tmp)
        temp_name = tmp.name
    shutil.move(temp_name, path)


def clear_pdf_metadata(path: str) -> None:
    reader = PdfReader(path)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        writer.write(tmp)
        temp_name = tmp.name
    shutil.move(temp_name, path)


class CurriculumApp(tk.Tk):
    def __init__(self) -> None:
        super().__init__()
        self.title("Gerenciador de Currículos")
        self.geometry("800x600")

        notebook = ttk.Notebook(self)
        notebook.pack(fill=tk.BOTH, expand=True)

        self.placeholder_entries: dict[str, tk.Entry] = {}
        self.template_path: str | None = None

        self.tab_template = ttk.Frame(notebook)
        notebook.add(self.tab_template, text="Preencher Currículo")
        self._build_template_tab()

        self.tab_clear = ttk.Frame(notebook)
        notebook.add(self.tab_clear, text="Limpar Metadados")
        self._build_clear_tab()

        self.tab_read = ttk.Frame(notebook)
        notebook.add(self.tab_read, text="Ler Metadados")
        self._build_read_tab()

        self.tab_edit = ttk.Frame(notebook)
        notebook.add(self.tab_edit, text="Editar Metadados")
        self._build_edit_tab()

    def _build_template_tab(self) -> None:
        frame = ttk.Frame(self.tab_template, padding=20)
        frame.pack(fill=tk.BOTH, expand=True)

        ttk.Button(frame, text="Selecionar Modelo (.docx)", command=self._load_template).pack(anchor=tk.W)

        self.template_label = ttk.Label(frame, text="Nenhum arquivo selecionado")
        self.template_label.pack(anchor=tk.W, pady=(10, 20))

        self.placeholder_frame = ttk.LabelFrame(frame, text="Variáveis Detectadas")
        self.placeholder_frame.pack(fill=tk.BOTH, expand=True)

        self.generate_button = ttk.Button(frame, text="Gerar Currículo", command=self._generate_document)
        self.generate_button.pack(anchor=tk.E, pady=10)

    def _load_template(self) -> None:
        path = filedialog.askopenfilename(title="Selecione o currículo modelo", filetypes=[("Documentos Word", "*.docx")])
        if not path:
            return
        try:
            placeholders = extract_placeholders(path)
        except Exception as exc:
            messagebox.showerror("Erro", f"Não foi possível ler o arquivo: {exc}")
            return

        self.template_path = path
        self.template_label.configure(text=f"Modelo: {os.path.basename(path)}")

        for widget in self.placeholder_frame.winfo_children():
            widget.destroy()
        self.placeholder_entries.clear()

        if not placeholders:
            ttk.Label(self.placeholder_frame, text="Nenhuma variável encontrada. Utilize o padrão {{VARIAVEL}}.").pack(anchor=tk.W)
        else:
            for placeholder in placeholders:
                row = ttk.Frame(self.placeholder_frame)
                row.pack(fill=tk.X, pady=2)
                ttk.Label(row, text=placeholder, width=25).pack(side=tk.LEFT)
                entry = ttk.Entry(row)
                entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
                self.placeholder_entries[placeholder] = entry

    def _generate_document(self) -> None:
        if not self.template_path:
            messagebox.showwarning("Atenção", "Selecione um modelo primeiro.")
            return
        replacements = {name: entry.get() for name, entry in self.placeholder_entries.items()}
        missing = [name for name, value in replacements.items() if not value]
        if missing:
            if not messagebox.askyesno("Campos vazios", "Algumas variáveis estão vazias. Deseja continuar?"):
                return
        save_path = filedialog.asksaveasfilename(
            title="Salvar currículo preenchido",
            defaultextension=".docx",
            filetypes=[("Documentos Word", "*.docx")],
        )
        if not save_path:
            return
        try:
            replace_placeholders(self.template_path, replacements, save_path)
            messagebox.showinfo("Sucesso", f"Currículo salvo em {save_path}")
        except Exception as exc:
            messagebox.showerror("Erro", f"Não foi possível gerar o currículo: {exc}")

    def _build_clear_tab(self) -> None:
        frame = ttk.Frame(self.tab_clear, padding=20)
        frame.pack(fill=tk.BOTH, expand=True)

        ttk.Button(frame, text="Selecionar arquivo", command=self._clear_metadata).pack(anchor=tk.W)
        ttk.Label(frame, text="O arquivo original será sobrescrito com os metadados removidos.").pack(anchor=tk.W, pady=10)

    def _clear_metadata(self) -> None:
        path = filedialog.askopenfilename(
            title="Selecione o currículo",
            filetypes=[("Documentos", "*.docx *.pdf"), ("Word", "*.docx"), ("PDF", "*.pdf")],
        )
        if not path:
            return
        try:
            if path.lower().endswith(".docx"):
                clear_docx_metadata(path)
            elif path.lower().endswith(".pdf"):
                clear_pdf_metadata(path)
            else:
                raise ValueError("Formato não suportado")
            messagebox.showinfo("Sucesso", "Metadados removidos com sucesso!")
        except Exception as exc:
            messagebox.showerror("Erro", f"Não foi possível limpar os metadados: {exc}")

    def _build_read_tab(self) -> None:
        frame = ttk.Frame(self.tab_read, padding=20)
        frame.pack(fill=tk.BOTH, expand=True)

        ttk.Button(frame, text="Selecionar arquivo", command=self._read_metadata).pack(anchor=tk.W)
        self.metadata_text = tk.Text(frame, height=20)
        self.metadata_text.pack(fill=tk.BOTH, expand=True, pady=10)
        self.metadata_text.configure(state=tk.DISABLED)

    def _read_metadata(self) -> None:
        path = filedialog.askopenfilename(
            title="Selecione o currículo",
            filetypes=[("Documentos", "*.docx *.pdf"), ("Word", "*.docx"), ("PDF", "*.pdf")],
        )
        if not path:
            return
        try:
            if path.lower().endswith(".docx"):
                data = read_docx_metadata(path)
            elif path.lower().endswith(".pdf"):
                data = read_pdf_metadata(path)
            else:
                raise ValueError("Formato não suportado")
        except Exception as exc:
            messagebox.showerror("Erro", f"Não foi possível ler os metadados: {exc}")
            return

        self.metadata_text.configure(state=tk.NORMAL)
        self.metadata_text.delete("1.0", tk.END)
        for key, value in data.items():
            self.metadata_text.insert(tk.END, f"{key}: {value}\n")
        self.metadata_text.configure(state=tk.DISABLED)

    def _build_edit_tab(self) -> None:
        frame = ttk.Frame(self.tab_edit, padding=20)
        frame.pack(fill=tk.BOTH, expand=True)

        ttk.Button(frame, text="Selecionar arquivo", command=self._load_for_edit).pack(anchor=tk.W)
        self.edit_file_label = ttk.Label(frame, text="Nenhum arquivo selecionado")
        self.edit_file_label.pack(anchor=tk.W, pady=(10, 20))

        form = ttk.Frame(frame)
        form.pack(fill=tk.X, pady=10)

        self.edit_entries: dict[str, tk.Entry] = {}
        fields = [
            ("DC: TITLE", "title"),
            ("DC: CREATOR", "creator"),
            ("CP: DESCRIPTION", "description"),
            ("CP: CATEGORY", "category"),
        ]
        for label_text, key in fields:
            row = ttk.Frame(form)
            row.pack(fill=tk.X, pady=5)
            ttk.Label(row, text=label_text, width=20).pack(side=tk.LEFT)
            entry = ttk.Entry(row)
            entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
            self.edit_entries[key] = entry

        ttk.Button(frame, text="Salvar Metadados", command=self._save_metadata).pack(anchor=tk.E, pady=10)
        self.editing_path: str | None = None

    def _load_for_edit(self) -> None:
        path = filedialog.askopenfilename(
            title="Selecione o currículo",
            filetypes=[("Documentos", "*.docx *.pdf"), ("Word", "*.docx"), ("PDF", "*.pdf")],
        )
        if not path:
            return
        try:
            if path.lower().endswith(".docx"):
                metadata = read_docx_metadata(path)
            elif path.lower().endswith(".pdf"):
                metadata = read_pdf_metadata(path)
            else:
                raise ValueError("Formato não suportado")
        except Exception as exc:
            messagebox.showerror("Erro", f"Não foi possível carregar os metadados: {exc}")
            return

        self.editing_path = path
        self.edit_file_label.configure(text=f"Arquivo: {os.path.basename(path)}")

        mapping = {
            "title": metadata.get("DC:TITLE", ""),
            "creator": metadata.get("DC:CREATOR", ""),
            "description": metadata.get("CP:DESCRIPTION", ""),
            "category": metadata.get("CP:CATEGORY", ""),
        }
        for key, entry in self.edit_entries.items():
            entry.delete(0, tk.END)
            entry.insert(0, mapping.get(key, ""))

    def _save_metadata(self) -> None:
        if not self.editing_path:
            messagebox.showwarning("Atenção", "Selecione um arquivo primeiro.")
            return
        title = self.edit_entries["title"].get()
        creator = self.edit_entries["creator"].get()
        description = self.edit_entries["description"].get()
        category = self.edit_entries["category"].get()
        try:
            if self.editing_path.lower().endswith(".docx"):
                write_docx_metadata(self.editing_path, title, creator, description, category)
            elif self.editing_path.lower().endswith(".pdf"):
                write_pdf_metadata(self.editing_path, title, creator, description, category)
            else:
                raise ValueError("Formato não suportado")
            messagebox.showinfo("Sucesso", "Metadados atualizados!")
        except Exception as exc:
            messagebox.showerror("Erro", f"Não foi possível salvar os metadados: {exc}")


def main() -> None:
    app = CurriculumApp()
    app.mainloop()


if __name__ == "__main__":
    main()
